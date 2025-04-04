import pandas as pd
import json
from docx import Document
import os

import pandas as pd
from docx import Document
import json
from transformers import (
    Gemma3ForConditionalGeneration,
    BitsAndBytesConfig,
    AutoProcessor,
)
import torch
import re
from openai import OpenAI
import pandas as pd

def chunk_text(text, max_chars=5000, overlap=10): #chunk text into smaller pieces to eventually feed through an llm. 
    # Split text into sentences (basic split using punctuation)
    sentences = re.split(r'(?<=[.!?]) +', text)
    
    chunks = []
    current_chunk = []
    char_count = 0
    
    i = 0
    while i < len(sentences):
        sentence = sentences[i]
        
        if char_count + len(sentence) <= max_chars or not current_chunk:
            current_chunk.append(sentence)
            char_count += len(sentence) + 1  # +1 for the space after sentence
            i += 1
        else:
            chunks.append(" ".join(current_chunk))
            # Move back by the overlap count to ensure overlap
            i = max(0, i - overlap)
            current_chunk = []
            char_count = 0
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks

# Example usage:
long_text = "Your long paragraph of text goes here..."
chunks = chunk_text(long_text)

for idx, chunk in enumerate(chunks):
    print(f"Chunk {idx+1}:\n{chunk}\n{'-'*40}")

def process_spreadsheet(input_file, json_output, doc_output_folder, header_names):
    # Step 1: Read the spreadsheet
    df = pd.read_json(input_file)

    # Initialize an empty list to store data for JSON
    json_data = []

    # Step 2: Loop through each row in the filtered dataframe
    for index, row in df.iterrows():
        # Initialize a dictionary to store the row data
        row_dict = {}
        non_nan_values = row[pd.notna(row)]

        # Check if the row has more non-NaN values than the header names
        if len(non_nan_values) > len(header_names):
            # If more columns than expected, concatenate all but the last two cells
            concatenated_value = " ".join(
                [str(cell) for cell in non_nan_values[:-2] if pd.notna(cell)]
            )  # Concatenate everything except last two
            row_dict[header_names[0]] = "reformatted_index_" + str(
                index
            )  # rename first colum (id)

            row_dict[header_names[1]] = (
                concatenated_value  # Assign concatenated value to the second header column
            )
            # Assign last two cells to the corresponding headers
            for i, column in enumerate(header_names[-2:], start=1):
                row_dict[column] = non_nan_values.iloc[-3 + i]
        elif len(non_nan_values) <= len(header_names):
            # If the row doesn't have more columns, simply assign values as is
            for column, value in row.items():
                if pd.notna(value):  # Exclude NaN values
                    row_dict[column] = value


        # Step 3: Create a Word document for this row
        doc = Document()
        doc.add_heading(f"Row {index + 1} Data", 0)
        for column, value in row_dict.items():
            if column == "Transcript":
                chunks = chunk_text(value, max_chars=5000, overlap=10) #10 sentence overlap
                sentences = str(value).split(". ")
                doc.add_paragraph(f"{column}")
                # Add each sentence as a new paragraph
                for sentence in sentences:
                    if sentence.strip():  # Avoid adding empty sentences
                        doc.add_paragraph(f"{sentence.strip()}.")
            else:
                doc.add_paragraph(f"{column}: {value}")
        row_dict['chunks_transcript'] = chunks
        # Append the row data to json_data list
        json_data.append(row_dict)


        # Save the Word document
        doc_filename = f"{doc_output_folder}/row_{index + 1}.docx"
        doc.save(doc_filename)

    # Step 4: Write the JSON output to a file
    with open(json_output, "w") as json_file:
        json.dump(json_data, json_file, indent=4)

    print(
        f"Processing complete. JSON output saved to {json_output}. Word documents saved in {doc_output_folder}."
    )
    return json_data


def setup_model(model_id="google/gemma-3-4b-it"):
    model = Gemma3ForConditionalGeneration.from_pretrained(
        model_id, temperature=0.2, device_map="auto"
    ).eval()
    processor = AutoProcessor.from_pretrained(model_id)
    return model, processor


def convert_transcript_to_script(transcript, model, processor):
    prompt = f"""
        I have a transcript of a conversation recorded in a doctor's office. 
        Please convert this transcript into a script format, making it easier to follow the conversation. 
        Do not change the wording, summarize, omit, or modify any part of the original transcript. 
        The only thing you should change is formatting, so it is structured like a script (e.g., speaker labels and clear separation between each speaker's lines).

        Here is the transcript:
        Transcript: {transcript}
    """
    messages = [
        {
            "role": "system",
            "content": [{"type": "text", "text": "You are a helpful assistant."}],
        },
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
            ],
        },
    ]
    input_text = processor.apply_chat_template(messages, add_generation_prompt=True)
    inputs = processor(
        images=None, text=input_text, add_special_tokens=False, return_tensors="pt"
    ).to(model.device)
    with torch.inference_mode():
        output = model.generate(**inputs, max_new_tokens=5000)
    input_len = inputs["input_ids"].shape[-1]
    output = output[0][input_len:]
    output = processor.decode(output, skip_special_tokens=True)
    print(output)

    return output


# Example usage
input_file = r"C:\Users\Liliana\dev\Ambience_testproject\omissions_sample_encounters.json"  # Input Excel file path
json_output = "output_data.json"  # Output JSON file path
doc_output_folder = r"C:\Users\Liliana\dev\Ambience_testproject\word_documents_split"  # Folder to save the Word documents
model_id="google/gemma-3-4b-it"

# Create the folder for Word documents if it doesn't exist
if not os.path.exists(doc_output_folder):
    os.makedirs(doc_output_folder)

# Call the function to process the spreadsheet
json_data = process_spreadsheet(
    input_file,
    json_output,
    doc_output_folder,
    header_names=["Encounter ID", "Transcript", "HPI", "Clinical feedback"],
)

# model, processor = setup_model(model_id=model_id)

# for ind, transcript in enumerate(json_data):
#     script = convert_transcript_to_script(transcript['Transcript'], model, processor) #use local models for now because I'm cheap
#     json_data[ind]['script'] = script

client = OpenAI()


for ind, encounter in enumerate(json_data):
        summarized_chunks = []
        for chunk in encounter['chunks_transcript']:
            query  = f"""Summarize everything in the transcript. Do not fabricate anything, or leave anything out. 
                Here is the transcript: {chunk}"""    
            
            completion = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[
                    {
                        "role": "user",
                        "content": query
                    }
                ]
            )
            print(completion.choices[0].message.content)
            summarized_chunks.append(completion.choices[0].message.content)
        json_data[ind]['summarized_chunks']=summarized_chunks

    

